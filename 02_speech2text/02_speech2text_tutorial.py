#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Jan 12 12:47:45 2018

@author: Howard Chung
"""

# Imports the Google Cloud client library
from google.cloud import speech
from google.cloud.speech import enums
from google.cloud.speech import types
from google.cloud import storage
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Cm
import datetime
import io
import os
import time
import jieba
jieba.set_dictionary('dict.txt.big.txt')
from wordcloud import WordCloud
from scipy.misc import imread
from os import path  
from nltk.tokenize import word_tokenize



# --------一分鐘內語音辨識 - 實作---------------
'''
# 語音基礎串接
一分鐘內語音辨識  (#不用經 bucket, 大於一分 要經 bucket )
'''
## 此練念, 其實只有二句

'''
* json_os：憑證檔的路徑
* title_pattern：錄音檔的名稱模式
* sample_rate_hertz：錄音的取樣頻率
* doc_title：docx文件名稱
* wd：工作目錄
* google bucket: speech2text_imphoenix	:
'''

#json_os = '憑證檔.json'
json_os = 'Speech2text憑證檔.json'

title_pattern='nlp'
sample_rate_hertz = 48000
doc_title = '範例1_一分鐘內雲端運算'
bucket = 'speech2text_imphoenix'

#os.getcwd() : 工作目錄 
#record 是 folder

wd = os.getcwd().split('02_speech2text')[0] + 'record' 


# 從python client端對雲端speech2text服務進行驗證
# **************************************************
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] =json_os
client = speech.SpeechClient()

file_list = os.listdir(wd)     

# 選出title_pattern的錄音檔
music = []
for i in file_list:
    if title_pattern in i:
        music.append(wd + '/' +i)

# music是含有2個路徑的list，假設我們要選擇list中裡面的第二個 nlp yes.wav 範例
type(music)
music

#  nlpno 是要捉出的 nlpno.wav 

music = [i for i in music if 'nlpno' in i] # mac 不一樣是，應該用pattern去做 請記得，引數都是從 0 開始，所以第 2 個數值就是引數 1 
music = [i for i in music if 'wav' in i] # mac 不一樣是，應該用pattern去做 請記得，引數都是從 0 開始，所以第 2 個數值就是引數 1 

##以上 會捉到  nlpno.wav 
music = music[0]# 成功抓出第二個數值



# 將 audio錄音檔 讀進來
with io.open(music, 'rb') as audio_file:
    content = audio_file.read()

# 看一下讀進python裡面是以何種方式呈現
content[0:100]

# 將錄音檔轉換成google 看得懂的格式
audio = types.RecognitionAudio(content=content)

# 如果您覺得自己電腦效能夠好的話，可以執行看看下面的audio
#audio

##################################################################
# google 語音分析
###################################################################

# 設定格式錄音檔
config = types.RecognitionConfig(
     encoding=enums.RecognitionConfig.AudioEncoding.LINEAR16, # LINEAR16 for Audio
    sample_rate_hertz=sample_rate_hertz,
    language_code='cmn-Hant-TW' ,  # for Taiwan 語系
    enable_word_time_offsets=True)   # 是否要分段

# 機器學習文字辨識(speech2text)
print('機器學習文字辨識')
response = client.recognize(config, audio)   #config file , audio for audio file
response

##################################################################
# data 整理, 把 語音轉成文字後
###################################################################

# 開始整理辨識出來的raw data，轉變成有用的資訊
'''
response -> result ( raw data的格式 ) -> alternatives ( raw data的格式 ) -> transcript (完整中文資料)
'''
#把 語音轉成文字後
response.results[0].alternatives[0].transcript

# 每一個result都是音檔一個個段落的辨識結果，所以我們用for迴圈的方式來得到【中文transcript】
for result in response.results:
    alternative = result.alternatives[0]
    print('Transcript: {}'.format(alternative.transcript))
    print('Confidence: {}'.format(alternative.confidence))
   
# raw data 格式    
alternative

##################################################################
# data 整理,處理 Time stamp
###################################################################

transcript_list = []
transcript_confidence = []
timerecored = []
# begining and end time of a sentence
sentence_start_time = alternative.words[0].start_time
sentence_end_time = alternative.words[len(alternative.words)-1].end_time

# make time
sentence_start_time = round( sentence_start_time.seconds + sentence_start_time.nanos * 1e-9)
sentence_end_time = round( sentence_end_time.seconds + sentence_end_time.nanos * 1e-9)

# make min
sentence_start_time= str(datetime.timedelta(seconds=sentence_start_time))
sentence_end_time =str(datetime.timedelta(seconds=sentence_end_time))
timerecored.append([sentence_start_time, sentence_end_time])


# 我們將transcript的結果儲存到不同的list裡面，留作後續資料整理之用途
'''
transcript_list : 存中文transcript
transcript_confidence : 存信心程度
timerecored : 存始末之時間
'''

transcript_list = []
transcript_confidence = []
timerecored = []

for result in response.results:
    alternative = result.alternatives[0]
    # The first alternative is the most likely one for this portion.
    transcript_list.append(alternative.transcript)
    transcript_confidence.append(alternative.confidence)
    print('Transcript: {}'.format(alternative.transcript))
    print('Confidence: {}'.format(alternative.confidence))
        
    # begining and end time of a sentence
    sentence_start_time = alternative.words[0].start_time
    sentence_end_time = alternative.words[len(alternative.words)-1].end_time
    
    # make time
    sentence_start_time = round( sentence_start_time.seconds + sentence_start_time.nanos * 1e-9)
    sentence_end_time = round( sentence_end_time.seconds + sentence_end_time.nanos * 1e-9)
    
    # make min
    sentence_start_time= str(datetime.timedelta(seconds=sentence_start_time))
    sentence_end_time =str(datetime.timedelta(seconds=sentence_end_time))
    timerecored.append([sentence_start_time, sentence_end_time])
    
    
# pandas 建立信心程度資料表
transcript_list
transcript_confidence
timerecored

# make df
transcript_df = pd.DataFrame(transcript_list, columns = ['文章段句'])
confidence_df = pd.DataFrame(transcript_confidence, columns = ['機器認字信心水準'])
confidence_df['機器認字信心水準'] = round(confidence_df['機器認字信心水準'],2)
time_df  = pd.DataFrame(timerecored, columns = ['start', 'end'])

#資料合在一起 pd.concat 

correctness_summary_df = pd.concat([transcript_df , confidence_df, time_df], axis = 1)    
correctness_summary_df = correctness_summary_df.sort_values(['機器認字信心水準'])
correctness_summary_df['改善順序'] = range(1, len(correctness_summary_df)+1)


timer_translist =[]
for hah,timer in zip(transcript_list,timerecored):
   timer_translist.append(hah+'  ' +'【'+' to '.join(timer)+'】')


# 製作文字雲
from speech2text import make_worldcould_report, text_freq
cut_text = make_worldcould_report(data = correctness_summary_df, pd_text_col = '文章段句',mask_pic = False, filename = 'wordcloud',pic_name = 'test.png')

##################################33
##  切詞

words_counts=text_freq(cut_text)

# 計算重要程度
max = words_counts['counts'].describe()['max']
mean = words_counts['counts'].describe()['mean']  

# 僅取出max與mean的字詞, 取 > mean 的字詞
words_counts = words_counts[ (words_counts['counts'] <= max) & (words_counts['counts']>=mean)  ]

###############################
# 重要字詞的加權 


df_count_all = pd.DataFrame()
for index, i in words_counts.iterrows():
    
    df_count = correctness_summary_df[correctness_summary_df['文章段句'].str.contains(i['word'])]
    
    if not df_count.empty :
        df_count['重要性'] = i['counts']   #重要性 = 出現次數 
        df_count_all  = pd.concat([df_count_all ,df_count ])
    
# group by    
correctness_summary_df = df_count_all.groupby(['文章段句', '機器認字信心水準', 'start', 'end', '改善順序'], as_index=False)['重要性'].mean().round(2)  #round(2) 四拾五入到小數點 2 位


# save to docx
document = Document()
document.add_heading(doc_title, 0)

#\n\n 空二行
document.add_paragraph('機器認字信心水準' + str(round(correctness_summary_df['機器認字信心水準'].mean(),2) ) + '\n\n' + '\n\n'.join(timer_translist) )

document.add_picture('wordcloud.png',width=Cm(15), height=Cm(13))
document.save(doc_title+'_文章逐字稿.docx')
correctness_summary_df.to_excel(doc_title+'_文章認字信心矩陣.xlsx') 
print('Done')
print('請看工作目錄檔案中有沒有兩個檔案，一格個是完整的docx檔案，一個是xlsx檔案')


# --------一分鐘內語音辨識 - function---------------
### 整合成所有的 function

# 從這邊開始執行
''' maked by phoenix 
def speech_to_text_in_a_min(doc_title = '範例1_一分鐘內雲端運算', title_pattern='nlpno', 
                            wd ='/home/slave1/git/Speech2Text_workshop/record',
                            json_os = '/home/slave1/git/Speech2Text_workshop/speech2text-3de4444fd46a.json',sample_rate_hertz =  48000):
 
'''    
def speech_to_text_in_a_min(doc_title = '範例1_一分鐘內雲端運算', title_pattern='nlpno.wav', 
                            wd ='/home/slave1/git/Speech2Text_workshop/record',
                            json_os = '/home/slave1/git/Speech2Text_workshop/speech2text-3de4444fd46a.json',sample_rate_hertz =  48000):
     
    '''
    * json_os：憑證檔的路徑
    * title_pattern：錄音檔的名稱模式
    * sample_rate_hertz：錄音的取樣頻率
    * doc_title：docx文件名稱
    * wd：工作目錄
    
    '''
    
    # 計時
    start_time = time.time()
    # 從python client端對雲端speech2text服務進行驗證
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] =json_os
    client = speech.SpeechClient()
    
    file_list = os.listdir(wd)     
    
    # 選出title_pattern的錄音檔
    select_wav = []
    for i in file_list:
        if title_pattern in i:
            select_wav.append(wd + '/' +i)
    
    aa = pd.DataFrame()
    
    for music in select_wav:
            
        # 將 audio錄音檔 讀入進來
        with io.open(music, 'rb') as audio_file:
            content = audio_file.read()
        
        # 將錄音檔轉換成google 看得懂的格式
        audio = types.RecognitionAudio(content=content)
        
        # 設定格式錄音檔
        config = types.RecognitionConfig(
             encoding=enums.RecognitionConfig.AudioEncoding.LINEAR16,
            sample_rate_hertz=sample_rate_hertz,
            language_code='cmn-Hant-TW' ,
            enable_word_time_offsets=True)
        
        # 機器學習文字辨識(speech2text)
        print('')
        response = client.recognize(config, audio)
            
        
        transcript_list = []
        transcript_confidence = []
        timerecored = []
        # Each result is for a consecutive portion of the audio. Iterate through
        # them to get the transcripts for the entire audio file.
        for result in response.results:
            alternative = result.alternatives[0]
            # The first alternative is the most likely one for this portion.
            transcript_list.append(alternative.transcript)
            transcript_confidence.append(alternative.confidence)
            print('Transcript: {}'.format(alternative.transcript))
            print('Confidence: {}'.format(alternative.confidence))
            
            
            # begining and end time of a sentence
            sentence_start_time = alternative.words[0].start_time
            sentence_end_time = alternative.words[len(alternative.words)-1].end_time
            
            # make time
            sentence_start_time = round( sentence_start_time.seconds + sentence_start_time.nanos * 1e-9)
            sentence_end_time = round( sentence_end_time.seconds + sentence_end_time.nanos * 1e-9)
            
            # make min
            sentence_start_time= str(datetime.timedelta(seconds=sentence_start_time))
            sentence_end_time =str(datetime.timedelta(seconds=sentence_end_time))
            timerecored.append([sentence_start_time, sentence_end_time])
            
        # pandas 建立信心程度資料表
         # make df
        transcript_df = pd.DataFrame(transcript_list, columns = ['文章段句'])
        confidence_df = pd.DataFrame(transcript_confidence, columns = ['機器認字信心水準'])
        confidence_df['機器認字信心水準'] = round(confidence_df['機器認字信心水準'],2)
        time_df  = pd.DataFrame(timerecored, columns = ['start', 'end'])
        correctness_summary_df = pd.concat([transcript_df , confidence_df,time_df], axis = 1)    
        correctness_summary_df = correctness_summary_df.sort_values(['機器認字信心水準'])
        correctness_summary_df['改善順序'] = range(1, len(correctness_summary_df)+1)
        
        timer_translist =[]
        for hah,timer in zip(transcript_list,timerecored):
           timer_translist.append(hah+'  ' +'【'+' to '.join(timer)+'】')
        
        aa = pd.concat([ aa, correctness_summary_df])
    
     
    
    # 製作文字雲
    from speech2text import make_worldcould_report, text_freq
    cut_text = make_worldcould_report(data = aa, pd_text_col = '文章段句',mask_pic = False, filename = 'wordcloud',pic_name = 'test.png')
    words_counts=text_freq(cut_text)
    
    # 計算重要程度
    max = words_counts['counts'].describe()['max']
    mean = words_counts['counts'].describe()['mean']
    
    # 僅取出max與mean的字詞
    words_counts = words_counts[ (words_counts['counts'] <= max) & (words_counts['counts']>=mean)  ]
    
    df_count_all = pd.DataFrame()
    for index, i in words_counts.iterrows():
        
        df_count = correctness_summary_df[correctness_summary_df['文章段句'].str.contains(i['word'])]
        
        if not df_count.empty :
            df_count['重要性'] = i['counts']
            df_count_all  = pd.concat([df_count_all ,df_count ])
        
    # group by    
    correctness_summary_df = df_count_all.groupby(['文章段句', '機器認字信心水準', 'start', 'end', '改善順序'], as_index=False)['重要性'].mean().round(2)


    # save to docx
    document = Document()
    document.add_heading(doc_title, 0)
    document.add_paragraph('機器認字信心水準' + str(round(correctness_summary_df['機器認字信心水準'].mean(),2) ) + '\n\n' + '\n\n'.join(timer_translist) )
    document.add_picture('wordcloud.png',width=Cm(15), height=Cm(13))
    document.save(doc_title+'_文章逐字稿.docx')
    print('Done')
    print('請看工作目錄檔案中有沒有兩個檔案，一格個是完整的docx檔案，一個是xlsx檔案')
    print("--- %s seconds ---" % (round(time.time() - start_time, 2)))
    return correctness_summary_df.to_excel(doc_title+'_文章認字信心矩陣.xlsx') 

# 請執行到這邊停住

# 一個function實戰搞定
speech_to_text_in_a_min(doc_title = '範例2_一分鐘內雲端運算', title_pattern='nlpno.wav', 
                            wd = os.getcwd().split('02_speech2text')[0] + 'record' ,
                            json_os = '憑證檔.json', sample_rate_hertz =  48000)


# --------雲端運算儲存空間 - function---------------

def save_to_gcp_storage(bucket = 'speech2textgood', title_pattern = 'nlpno.wav', wd ='/home/slave1/git/Speech2Text_workshop/record'):
    # list out files
    file_list = os.listdir(wd)     
    file_list = list(filter(lambda x:title_pattern in x, file_list))
            
    # store init
    storage_client = storage.Client()
    bucket = storage_client.get_bucket(bucket)
    
    # for loop to store
    for i in file_list:
        blob = bucket.blob(i)
        blob.upload_from_filename(wd + '/' + i)
        print('已經成功上傳' + i  +'到'+str(bucket)+'的雲端運算儲存空間')
    
    return file_list

# --------一分鐘以上語音辨識 - 實作---------------
# 一個function實戰搞定 - 看能否work?
speech_to_text_in_a_min(doc_title = '範例1_一分鐘內雲端運算', title_pattern='sent3.wav', 
                            wd = os.getcwd().split('02_speech2text')[0] + 'record' ,
                            json_os = 'Speech2text憑證檔.json', sample_rate_hertz =  48000)

## error: 400 s
# InvalidArgument: 400 Sync input too long. For audio longer than 1 min use LongRunningRecognize with a 'uri' parameter.

# 首先要先定義要上傳雲端空間的檔案並且利用程式上傳上去
wd =  os.getcwd().split('02_speech2text')[0] + 'record' 

##
#bucket = 'speech2textgood'
bucket = 'speech2text_imphoenix'


name = 'sent3.wav'
file_list = save_to_gcp_storage(bucket = bucket, title_pattern = name, wd =wd )

# 雲端運算的參數

# gcs_uri 可以用兩種形式表示
gcs_uri = 'gs://speech2textgood/sent3.wav'
gcs_uri 

gcs_uri = 'gs://'+bucket +'/' +name
gcs_uri 

timeout = None

# 從python client端對雲端speech2text服務進行驗證
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] =json_os

client = speech.SpeechClient()

audio = types.RecognitionAudio(uri=gcs_uri)
audio

config = types.RecognitionConfig(
    encoding=enums.RecognitionConfig.AudioEncoding.LINEAR16,
    sample_rate_hertz=sample_rate_hertz,
    language_code='cmn-Hant-TW' ,
    enable_word_time_offsets=True)

# 機器學習文字辨識(speech2text)
print('雲端運算 - 機器學習辨識中')
operation = client.long_running_recognize(config, audio)
response = operation.result(timeout=timeout) 


aa = pd.DataFrame()
transcript_list = []
transcript_confidence = []
timerecored =[]
# Each result is for a consecutive portion of the audio. Iterate through
# them to get the transcripts for the entire audio file.
for result in response.results:
    alternative = result.alternatives[0]
    # The first alternative is the most likely one for this portion.
    transcript_list.append(alternative.transcript)
    transcript_confidence.append(alternative.confidence)
    print('Transcript: {}'.format(alternative.transcript))
    print('Confidence: {}'.format(alternative.confidence))
    
    # begining and end time of a sentence
    sentence_start_time = alternative.words[0].start_time
    sentence_end_time = alternative.words[len(alternative.words)-1].end_time
    
    # make time
    sentence_start_time = round( sentence_start_time.seconds + sentence_start_time.nanos * 1e-9)
    sentence_end_time = round( sentence_end_time.seconds + sentence_end_time.nanos * 1e-9)
    
    # make min
    sentence_start_time= str(datetime.timedelta(seconds=sentence_start_time))
    sentence_end_time =str(datetime.timedelta(seconds=sentence_end_time))
    timerecored.append([sentence_start_time, sentence_end_time])
    
    
    # pandas 建立信心程度資料表
     # make df
    transcript_df = pd.DataFrame(transcript_list, columns = ['文章段句'])
    confidence_df = pd.DataFrame(transcript_confidence, columns = ['機器認字信心水準'])
    confidence_df['機器認字信心水準'] = round(confidence_df['機器認字信心水準'],2)
    time_df  = pd.DataFrame(timerecored, columns = ['start', 'end'])
    correctness_summary_df = pd.concat([transcript_df , confidence_df,time_df], axis = 1)    
    correctness_summary_df = correctness_summary_df.sort_values(['機器認字信心水準'])
    correctness_summary_df['改善順序'] = range(1, len(correctness_summary_df)+1)
    
    timer_translist =[]
    for hah,timer in zip(transcript_list,timerecored):
       timer_translist.append(hah+'  ' +'【'+' to '.join(timer)+'】')
    aa = pd.concat([ aa, correctness_summary_df])


# 製作文字雲
from speech2text import make_worldcould_report, text_freq
cut_text = make_worldcould_report(data = aa, pd_text_col = '文章段句',mask_pic = False, filename = 'wordcloud',pic_name = 'test.png')
words_counts=text_freq(cut_text)

# 計算重要程度
max = words_counts['counts'].describe()['max']
mean = words_counts['counts'].describe()['mean']

# 僅取出max與mean的字詞
words_counts = words_counts[ (words_counts['counts'] <= max) & (words_counts['counts']>=mean)  ]

df_count_all = pd.DataFrame()
for index, i in words_counts.iterrows():
    
    df_count = correctness_summary_df[correctness_summary_df['文章段句'].str.contains(i['word'])]
    
    if not df_count.empty :
        df_count['重要性'] = i['counts']
        df_count_all  = pd.concat([df_count_all ,df_count ])
    
# group by    
correctness_summary_df = df_count_all.groupby(['文章段句', '機器認字信心水準', 'start', 'end', '改善順序'], as_index=False)['重要性'].mean().round(2)


# save to docx
document = Document()
document.add_heading(doc_title, 0)
document.add_paragraph('機器認字信心水準' + str(round(correctness_summary_df['機器認字信心水準'].mean(),2) ) + '\n\n' + '\n\n'.join(timer_translist) )
document.add_picture('wordcloud.png',width=Cm(15), height=Cm(13))
document.save(doc_title+'_文章逐字稿.docx')
correctness_summary_df.to_excel(doc_title+'_文章認字信心矩陣.xlsx') 
print('Done')
print('請看工作目錄檔案中有沒有兩個檔案，一格個是完整的docx檔案，一個是xlsx檔案')

#################################################3
#
#


# --------一分鐘以上語音辨識 - function---------------

# 從這邊開始執行
def speech_to_text(gcs_uri = 'gs://speechfashion/Acc.wav', doc_title = '範例2_一分鐘以上雲端運算', timeout = None,
                   json_os = '/home/slave1/git/Speech2Text/damnhow-db8d83229dd4.json',sample_rate_hertz =  96000):
    '''
    1.產出文章認字信心矩陣csv，提供修改者文句之修正順序
    2.產出docx文本，並提供文章機器認字信心水準，供修改者修改
    
    '''
     # 計時
    start_time = time.time()
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] =json_os
    
    client = speech.SpeechClient()

    audio = types.RecognitionAudio(uri=gcs_uri)
    
    config = types.RecognitionConfig(
        encoding=enums.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=sample_rate_hertz,
        language_code='cmn-Hant-TW' ,
        enable_word_time_offsets=True)
    
        #    config=types.StreamingRecognitionConfig(config=config)
        #    stream = [audio]
        #    requests = (types.StreamingRecognizeRequest(audio_content=chunk)
        #                for chunk in stream)
        #    responses = client.streaming_recognize(config, requests)

    operation = client.long_running_recognize(config, audio)

    print('機器學習文字辨識中...')
    response = operation.result(timeout=timeout) #
    
    aa = pd.DataFrame()
    transcript_list = []
    transcript_confidence = []
    timerecored =[]
    # Each result is for a consecutive portion of the audio. Iterate through
    # them to get the transcripts for the entire audio file.
    for result in response.results:
        alternative = result.alternatives[0]
        # The first alternative is the most likely one for this portion.
        transcript_list.append(alternative.transcript)
        transcript_confidence.append(alternative.confidence)
        print('Transcript: {}'.format(alternative.transcript))
        print('Confidence: {}'.format(alternative.confidence))
        
        # begining and end time of a sentence
        sentence_start_time = alternative.words[0].start_time
        sentence_end_time = alternative.words[len(alternative.words)-1].end_time
        
        # make time
        sentence_start_time = round( sentence_start_time.seconds + sentence_start_time.nanos * 1e-9)
        sentence_end_time = round( sentence_end_time.seconds + sentence_end_time.nanos * 1e-9)
        
        # make min
        sentence_start_time= str(datetime.timedelta(seconds=sentence_start_time))
        sentence_end_time =str(datetime.timedelta(seconds=sentence_end_time))
        timerecored.append([sentence_start_time, sentence_end_time])
        
        
        # pandas 建立信心程度資料表
         # make df
        transcript_df = pd.DataFrame(transcript_list, columns = ['文章段句'])
        confidence_df = pd.DataFrame(transcript_confidence, columns = ['機器認字信心水準'])
        confidence_df['機器認字信心水準'] = round(confidence_df['機器認字信心水準'],2)
        time_df  = pd.DataFrame(timerecored, columns = ['start', 'end'])
        correctness_summary_df = pd.concat([transcript_df , confidence_df,time_df], axis = 1)    
        correctness_summary_df = correctness_summary_df.sort_values(['機器認字信心水準'])
        correctness_summary_df['改善順序'] = range(1, len(correctness_summary_df)+1)
        
        timer_translist =[]
        for hah,timer in zip(transcript_list,timerecored):
           timer_translist.append(hah+'  ' +'【'+' to '.join(timer)+'】')
        
        aa = pd.concat([ aa, correctness_summary_df])
    
     
    
    # 製作文字雲
    from speech2text import make_worldcould_report, text_freq
    cut_text = make_worldcould_report(data = aa, pd_text_col = '文章段句',mask_pic = False, filename = 'wordcloud',pic_name = 'test.png')
    words_counts=text_freq(cut_text)
    
    # 計算重要程度
    max = words_counts['counts'].describe()['max']
    mean = words_counts['counts'].describe()['mean']
    
    # 僅取出max與mean的字詞
    words_counts = words_counts[ (words_counts['counts'] <= max) & (words_counts['counts']>=mean)  ]
    
    df_count_all = pd.DataFrame()
    for index, i in words_counts.iterrows():
        
        df_count = correctness_summary_df[correctness_summary_df['文章段句'].str.contains(i['word'])]
        
        if not df_count.empty :
            df_count['重要性'] = i['counts']
            df_count_all  = pd.concat([df_count_all ,df_count ])
        
    # group by    
    correctness_summary_df = df_count_all.groupby(['文章段句', '機器認字信心水準', 'start', 'end', '改善順序'], as_index=False)['重要性'].mean().round(2)


    # save to docx
    document = Document()
    document.add_heading(doc_title, 0)
    document.add_paragraph('機器認字信心水準' + str(round(correctness_summary_df['機器認字信心水準'].mean(),2) ) + '\n\n' + '\n\n'.join(timer_translist) )
    document.add_picture('wordcloud.png',width=Cm(15), height=Cm(13))
    document.save(doc_title+'_文章逐字稿.docx')
    print('Done')
    print('請看工作目錄檔案中有沒有兩個檔案，一格個是完整的docx檔案，一個是xlsx檔案')
    print("--- %s seconds ---" % (round(time.time() - start_time, 2)))
    return correctness_summary_df.to_excel(doc_title+'_文章認字信心矩陣.xlsx') 

# 請執行到這邊停住

# 一個function實戰搞定
    # gcs_uri = 'gs://speechfashion/Acc.wav'
speech_to_text(gcs_uri = 'gs://' +bucket+ '/'+ name, doc_title = '範例3_一分鐘以上雲端運算', timeout = None,
                   json_os =json_os, sample_rate_hertz =  48000)
